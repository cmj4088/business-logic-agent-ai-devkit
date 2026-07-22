"""ipd-docx Skill — IPD 项目 Word 文档生成。

支持 MRD、PRD、技术方案、测试报告、商业论证等 IPD 正式交付物。
生成格式规范的 .docx 文件，含样式、目录、表格、页眉页脚。
"""

import io
import json
import logging
from datetime import datetime
from typing import Any
from .base import BaseSkill, SkillContext, SkillResult

logger = logging.getLogger(__name__)

# 文档模板类型定义
DOCUMENT_TYPES = {
    "mrd": {
        "name": "市场需求文档",
        "description": "MRD：执行摘要、市场背景、用户画像、竞品分析、需求清单、产品定位",
        "sections": [
            "执行摘要", "市场背景与趋势", "目标用户画像",
            "竞品分析", "市场需求清单", "产品定位", "附录",
        ],
    },
    "prd": {
        "name": "产品需求文档",
        "description": "PRD：执行摘要、产品背景、功能需求、非功能需求、用户故事、验收标准",
        "sections": [
            "执行摘要", "产品背景与目标", "功能需求列表",
            "非功能需求", "用户故事", "验收标准", "附录",
        ],
    },
    "tech_spec": {
        "name": "技术方案文档",
        "description": "技术方案：系统架构、技术选型、模块设计、接口定义、风险评估",
        "sections": [
            "执行摘要", "系统架构", "技术选型",
            "模块设计", "接口定义", "风险评估", "附录",
        ],
    },
    "test_report": {
        "name": "测试报告",
        "description": "测试报告：测试概览、范围、环境、结果统计、缺陷分析、质量评估",
        "sections": [
            "测试概览", "测试范围", "测试环境",
            "测试结果统计", "缺陷分析", "质量评估", "附录",
        ],
    },
    "business_case": {
        "name": "商业论证",
        "description": "商业论证：市场机会、财务预测、风险评估、投资回报分析、建议",
        "sections": [
            "执行摘要", "市场机会", "财务预测",
            "风险评估", "投资回报分析", "建议", "附录",
        ],
    },
}


class DocxSkill(BaseSkill):
    """IPD Word 文档生成 Skill。

    支持 5 类文档模板：MRD、PRD、技术方案、测试报告、商业论证。
    生成格式规范的 .docx 文件。
    """

    @property
    def name(self) -> str:
        return "ipd-docx"

    @property
    def description(self) -> str:
        return "IPD 项目 Word 交付物生成：MRD、PRD、技术方案、测试报告、商业论证"

    def get_tools(self) -> list[dict]:
        """注册为 Agent 可调用的 tool。"""
        return [
            {
                "tool_name": "generate_docx",
                "tool_schema": json.dumps({
                    "type": "function",
                    "function": {
                        "name": "generate_docx",
                        "description": "生成 IPD 项目 Word 文档，支持 MRD/PRD/技术方案/测试报告/商业论证模板",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "document_type": {
                                    "type": "string",
                                    "enum": list(DOCUMENT_TYPES.keys()),
                                    "description": "文档类型",
                                },
                                "project_name": {
                                    "type": "string",
                                    "description": "项目名称",
                                },
                                "sections_content": {
                                    "type": "string",
                                    "description": "各章节内容（JSON 对象，key 为章节名，value 为 Markdown 内容）",
                                },
                                "author": {
                                    "type": "string",
                                    "description": "文档作者/创建者",
                                },
                                "version": {
                                    "type": "string",
                                    "description": "文档版本号，默认 v1.0",
                                },
                                "extra_params": {
                                    "type": "object",
                                    "description": "额外参数（如公司名、部门等）",
                                    "properties": {
                                        "company": {"type": "string"},
                                        "department": {"type": "string"},
                                        "confidential": {"type": "boolean", "description": "是否机密文档"},
                                    },
                                },
                            },
                            "required": ["document_type", "project_name"],
                        },
                    },
                }, ensure_ascii=False),
            },
        ]

    async def validate(self, context: SkillContext) -> tuple[bool, str]:
        params = context.params
        doc_type = params.get("document_type", "")
        if doc_type and doc_type not in DOCUMENT_TYPES:
            return False, f"不支持的文档类型 '{doc_type}'，可选: {list(DOCUMENT_TYPES.keys())}"
        return True, ""

    async def execute(self, context: SkillContext) -> SkillResult:
        params = context.params
        doc_type = params.get("document_type", "mrd")
        project_name = params.get("project_name", "未命名项目")
        sections_content = params.get("sections_content", "{}")
        author = params.get("author", "Business Logic Agent")
        version = params.get("version", "v1.0")
        extra_params = params.get("extra_params", {})

        doc_info = DOCUMENT_TYPES.get(doc_type, DOCUMENT_TYPES["mrd"])

        # 解析章节内容
        try:
            sections = json.loads(sections_content) if isinstance(sections_content, str) else sections_content
        except json.JSONDecodeError:
            sections = {}

        file_content = await self._generate_docx(
            doc_type=doc_type,
            doc_name=doc_info["name"],
            sections=doc_info["sections"],
            project_name=project_name,
            sections_content=sections,
            author=author,
            version=version,
            extra_params=extra_params,
        )

        # 保存文件到输出目录
        file_path = self._save_file(file_content, doc_type, project_name, version)

        return SkillResult(
            success=True,
            skill_name=self.name,
            output=f"Word 文档已生成: {doc_info['name']} (v{version})",
            file_path=file_path,
            file_type="docx",
            artifact_type=f"{doc_type}_docx",
            tokens_used=len(sections) * 50 + 200,
            metadata={
                "document_type": doc_type,
                "document_name": doc_info["name"],
                "sections": doc_info["sections"],
                "project_name": project_name,
                "version": version,
                "author": author,
            },
        )

    def _save_file(self, content: bytes, doc_type: str, project_name: str, version: str) -> str:
        """保存生成的文档到输出目录。"""
        import os
        output_dir = os.path.join(os.path.dirname(__file__), "..", "output", "skills")
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in project_name if c.isalnum() or c in "_ ").strip().replace(" ", "_")
        filename = f"{doc_type}_{safe_name}_{version}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        with open(filepath, "wb") as f:
            f.write(content)
        return filepath

    async def _generate_docx(
        self,
        doc_type: str,
        doc_name: str,
        sections: list[str],
        project_name: str,
        sections_content: dict,
        author: str,
        version: str,
        extra_params: dict,
    ) -> bytes:
        """生成 Word 文档内容。

        使用 python-docx 创建格式化的 .docx 文件。
        如果 python-docx 不可用，返回 Markdown 格式降级内容。

        Returns:
            bytes: Word 文件二进制内容。
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("python-docx 未安装，返回 Markdown 降级内容")
            return self._markdown_fallback(doc_name, sections, sections_content, project_name, author, version)

        doc = Document()

        # 设置默认样式
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.line_spacing = 1.5

        # ===== 封面 =====
        for _ in range(6):
            doc.add_paragraph("")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(doc_name)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        doc.add_paragraph("")

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"项目名称: {project_name}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

        doc.add_paragraph("")
        doc.add_paragraph("")

        # 封面信息
        info_items = [
            ("文档版本", version),
            ("创建日期", self._now()),
            ("作 者", author),
            ("文档类型", doc_name),
        ]
        for label, value in info_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}: {value}")
            run.font.size = Pt(11)

        # 分页
        doc.add_page_break()

        # ===== 目录页 =====
        doc.add_heading("目录", level=1)
        for section_name in sections:
            p = doc.add_paragraph()
            run = p.add_run(section_name)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

        doc.add_page_break()

        # ===== 正文 =====
        for section_idx, section_name in enumerate(sections, 1):
            doc.add_heading(f"{section_idx}. {section_name}", level=1)

            content = sections_content.get(section_name, "")
            if content:
                self._add_markdown_to_docx(doc, content)
            else:
                p = doc.add_paragraph()
                run = p.add_run("（此处填写内容）")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
                run.font.italic = True

            doc.add_paragraph("")

        # ===== 页脚：机密声明 =====
        if extra_params.get("confidential", False):
            footer = doc.sections[0].footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("机密 — 仅限内部使用")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # 保存到内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _add_markdown_to_docx(self, doc, markdown_text: str):
        """将 Markdown 格式文本添加到 docx 文档中。"""
        import re

        for line in markdown_text.split("\n"):
            line = line.strip()
            if not line:
                continue

            # 标题
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            # 列表
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            # 数字列表
            elif re.match(r"^\d+\.\s", line):
                p = doc.add_paragraph(line, style="List Number")
            # 表格
            elif "|" in line and line.count("|") >= 3:
                parts = [p.strip() for p in line.split("|") if p.strip()]
                if parts:
                    p = doc.add_paragraph(" | ".join(parts))
            else:
                doc.add_paragraph(line)

    def _markdown_fallback(
        self, doc_name: str, sections: list[str],
        sections_content: dict, project_name: str,
        author: str, version: str,
    ) -> bytes:
        """python-docx 不可用时，返回 Markdown 格式降级内容。"""
        lines = []
        lines.append(f"# {doc_name}")
        lines.append(f"")
        lines.append(f"> **项目**: {project_name}")
        lines.append(f"> **版本**: {version}")
        lines.append(f"> **作者**: {author}")
        lines.append(f"> **生成时间**: {self._now()}")
        lines.append(f"")
        lines.append(f"---")
        lines.append(f"")

        for section_idx, section_name in enumerate(sections, 1):
            lines.append(f"## {section_idx}. {section_name}")
            lines.append(f"")
            content = sections_content.get(section_name, "")
            if content:
                lines.append(content)
            else:
                lines.append(">（此处填写内容）")
            lines.append(f"")

        lines.append(f"---")
        lines.append(f"*文档由 `ipd-docx` Skill 自动生成*")

        return "\n".join(lines).encode("utf-8")

    def _now(self) -> str:
        from datetime import datetime, timezone
        return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")